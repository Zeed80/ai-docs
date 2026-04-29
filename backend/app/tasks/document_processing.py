from __future__ import annotations

import json
import base64
from io import BytesIO
from pathlib import Path

from backend.app.ai import AIRouter
from backend.app.ai.schemas import AIRequest, AITask, ChatMessage
from backend.app.domain.models import Document, ProcessingJobStatus
from backend.app.domain.schemas import (
    DocumentArtifactRead,
    DocumentExtractionResult,
    StructuredDocumentExtraction,
)
from backend.app.domain.storage import LocalFileStorage


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
CAD_EXTENSIONS = {".dxf", ".dwg", ".step", ".stp", ".iges", ".igs"}
TEXT_PREVIEW_LIMIT = 12000


class RawExtractionResult:
    def __init__(
        self,
        *,
        status: ProcessingJobStatus,
        parser_name: str,
        text: str = "",
        unsupported_reason: str | None = None,
    ) -> None:
        self.status = status
        self.parser_name = parser_name
        self.text = text
        self.unsupported_reason = unsupported_reason


async def process_document(
    document: Document,
    ai_router: AIRouter,
    storage: LocalFileStorage | None = None,
) -> DocumentExtractionResult:
    raw = extract_text(document)
    if raw.status != ProcessingJobStatus.COMPLETED:
        artifacts = create_preview_artifacts(document, storage) if storage else []
        ocr_text = await ocr_preview_artifacts(document, artifacts, ai_router)
        if ocr_text.strip():
            structured = await structure_extracted_text(document, ocr_text, ai_router)
            return DocumentExtractionResult(
                status=ProcessingJobStatus.COMPLETED.value,
                parser_name=f"{raw.parser_name}+ocr",
                text_preview=ocr_text[:TEXT_PREVIEW_LIMIT],
                text_length=len(ocr_text),
                unsupported_reason=raw.unsupported_reason,
                artifacts=artifacts,
                structured=structured,
            )
        return DocumentExtractionResult(
            status=raw.status.value,
            parser_name=raw.parser_name,
            text_preview=None,
            text_length=0,
            unsupported_reason=raw.unsupported_reason,
            artifacts=artifacts,
            structured=None,
        )

    structured = await structure_extracted_text(document, raw.text, ai_router)
    return DocumentExtractionResult(
        status=ProcessingJobStatus.COMPLETED.value,
        parser_name=raw.parser_name,
        text_preview=raw.text[:TEXT_PREVIEW_LIMIT],
        text_length=len(raw.text),
        artifacts=[],
        structured=structured,
    )


def extract_text(document: Document) -> RawExtractionResult:
    path = Path(document.storage_path)
    suffix = path.suffix.lower()
    if not path.exists():
        return RawExtractionResult(
            status=ProcessingJobStatus.FAILED,
            parser_name="missing_file",
            unsupported_reason="Stored file does not exist.",
        )
    if suffix in TEXT_EXTENSIONS:
        return _extract_plain_text(path)
    if suffix in DOCX_EXTENSIONS:
        return _extract_docx(path)
    if suffix in XLSX_EXTENSIONS:
        return _extract_xlsx(path)
    if suffix == ".pdf":
        return _extract_pdf_text_layer(path)
    if suffix in IMAGE_EXTENSIONS:
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="image_placeholder",
            unsupported_reason="Image OCR is not implemented yet; route to AI vision/OCR later.",
        )
    if suffix == ".dxf":
        return _extract_dxf(path)
    if suffix in {".step", ".stp"}:
        return _extract_step(path)
    if suffix in {".dwg", ".iges", ".igs"}:
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="cad_placeholder",
            unsupported_reason=(
                "This CAD format requires a dedicated geometry backend; DWG/IGES parsing is not "
                "enabled in the local Ubuntu pipeline."
            ),
        )
    return RawExtractionResult(
        status=ProcessingJobStatus.UNSUPPORTED,
        parser_name="unsupported_extension",
        unsupported_reason=f"Unsupported document extension: {suffix or 'none'}.",
    )


async def structure_extracted_text(
    document: Document,
    text: str,
    ai_router: AIRouter,
) -> StructuredDocumentExtraction:
    prompt = (
        "Extract verifiable manufacturing document metadata as JSON matching this shape: "
        "{document_type: string, summary: string, fields: "
        "[{name: string, value: string|number|boolean|null, confidence: number, reason: string, "
        "source: string|null}]}.\n"
        "Use confidence 0..1 and short field-level reasons. "
        "Do not propose actions or external tool calls.\n\n"
        f"Filename: {document.filename}\n"
        f"Content-Type: {document.content_type or 'unknown'}\n"
        f"Known type: {document.document_type or 'unknown'}\n"
        f"Text:\n{text[:TEXT_PREVIEW_LIMIT]}"
    )
    response = await ai_router.run(
        AIRequest(
            task=AITask.STRUCTURED_EXTRACTION,
            messages=[ChatMessage(role="user", content=prompt)],
            response_schema=StructuredDocumentExtraction,
            confidential=True,
            metadata={"document_id": document.id, "local_only": True},
        )
    )
    if isinstance(response.data, StructuredDocumentExtraction):
        return response.data
    if isinstance(response.data, dict):
        return StructuredDocumentExtraction.model_validate(response.data)
    return StructuredDocumentExtraction.model_validate(json.loads(response.text or "{}"))


def create_preview_artifacts(
    document: Document,
    storage: LocalFileStorage,
) -> list[DocumentArtifactRead]:
    path = Path(document.storage_path)
    suffix = path.suffix.lower()
    if not path.exists():
        return []
    if suffix in IMAGE_EXTENSIONS:
        return _create_image_preview(document, path, storage)
    if suffix == ".pdf":
        return _create_pdf_page_previews(document, path, storage)
    return []


async def ocr_preview_artifacts(
    document: Document,
    artifacts: list[DocumentArtifactRead],
    ai_router: AIRouter,
) -> str:
    image_artifacts = [
        artifact
        for artifact in artifacts
        if artifact.content_type and artifact.content_type.startswith("image/")
    ]
    if not image_artifacts:
        return ""
    images = [_artifact_as_data_uri(artifact) for artifact in image_artifacts]
    response = await ai_router.run(
        AIRequest(
            task=AITask.INVOICE_OCR,
            messages=[
                ChatMessage(
                    role="user",
                    content=(
                        "Transcribe visible text from these document preview images. "
                        "Return only the text and preserve important numbers, dates, supplier names, "
                        "drawing labels, and table rows."
                    ),
                )
            ],
            images=images,
            confidential=True,
            metadata={"document_id": document.id, "local_only": True},
        )
    )
    return response.text or ""


def _extract_plain_text(path: Path) -> RawExtractionResult:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        return RawExtractionResult(
            status=ProcessingJobStatus.FAILED,
            parser_name="plain_text",
            unsupported_reason=str(exc),
        )
    return RawExtractionResult(
        status=ProcessingJobStatus.COMPLETED,
        parser_name=f"text{path.suffix.lower()}",
        text=text,
    )


def _extract_pdf_text_layer(path: Path) -> RawExtractionResult:
    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError:
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="pdf_text_layer_unavailable",
            unsupported_reason="PyMuPDF is not installed; PDF text extraction is skipped safely.",
        )

    try:
        parts: list[str] = []
        with fitz.open(path) as pdf:
            for page in pdf:
                page_text = page.get_text("text").strip()
                if page_text:
                    parts.append(page_text)
        text = "\n\n".join(parts)
    except Exception as exc:
        return RawExtractionResult(
            status=ProcessingJobStatus.FAILED,
            parser_name="pdf_text_layer",
            unsupported_reason=str(exc),
        )

    if not text.strip():
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="pdf_text_layer",
            unsupported_reason="PDF text layer is empty; OCR/render fallback is not implemented yet.",
        )
    return RawExtractionResult(
        status=ProcessingJobStatus.COMPLETED,
        parser_name="pdf_text_layer",
        text=text,
    )


def _extract_docx(path: Path) -> RawExtractionResult:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError:
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="docx_unavailable",
            unsupported_reason="python-docx is not installed; DOCX text extraction is skipped safely.",
        )

    try:
        document = docx.Document(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        table_rows: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
        text = "\n".join([part for part in paragraphs + table_rows if part])
    except Exception as exc:
        return RawExtractionResult(
            status=ProcessingJobStatus.FAILED,
            parser_name="docx",
            unsupported_reason=str(exc),
        )

    if not text.strip():
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="docx",
            unsupported_reason="DOCX contains no extractable text.",
        )
    return RawExtractionResult(status=ProcessingJobStatus.COMPLETED, parser_name="docx", text=text)


def _extract_xlsx(path: Path) -> RawExtractionResult:
    try:
        import openpyxl  # type: ignore[import-not-found]
    except ImportError:
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="xlsx_unavailable",
            unsupported_reason="openpyxl is not installed; XLSX text extraction is skipped safely.",
        )

    try:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    lines.append(" | ".join(values))
        workbook.close()
        text = "\n".join(lines)
    except Exception as exc:
        return RawExtractionResult(
            status=ProcessingJobStatus.FAILED,
            parser_name="xlsx",
            unsupported_reason=str(exc),
        )

    if not text.strip():
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="xlsx",
            unsupported_reason="XLSX contains no extractable cell values.",
        )
    return RawExtractionResult(status=ProcessingJobStatus.COMPLETED, parser_name="xlsx", text=text)


def _extract_dxf(path: Path) -> RawExtractionResult:
    try:
        import ezdxf  # type: ignore[import-not-found]
    except ImportError:
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="dxf_unavailable",
            unsupported_reason="ezdxf is not installed; DXF extraction is skipped safely.",
        )

    try:
        doc = ezdxf.readfile(path)
        modelspace = doc.modelspace()
        entity_counts: dict[str, int] = {}
        layers = sorted({entity.dxf.layer for entity in modelspace if hasattr(entity.dxf, "layer")})
        for entity in modelspace:
            entity_type = entity.dxftype()
            entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
        text = (
            f"DXF version: {doc.dxfversion}\n"
            f"Layers: {', '.join(layers) if layers else 'none'}\n"
            f"Entity counts: {json.dumps(entity_counts, ensure_ascii=False, sort_keys=True)}"
        )
    except Exception as exc:
        return RawExtractionResult(
            status=ProcessingJobStatus.FAILED,
            parser_name="dxf",
            unsupported_reason=str(exc),
        )

    return RawExtractionResult(status=ProcessingJobStatus.COMPLETED, parser_name="dxf", text=text)


def _extract_step(path: Path) -> RawExtractionResult:
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        return RawExtractionResult(
            status=ProcessingJobStatus.FAILED,
            parser_name="step_header",
            unsupported_reason=str(exc),
        )

    header = _step_header(content)
    entity_counts = _step_entity_counts(content)
    text = (
        "STEP header preview. Full geometry analysis requires FreeCAD/pythonOCC backend.\n"
        f"{header}\n"
        f"Entity counts: {json.dumps(entity_counts, ensure_ascii=False, sort_keys=True)}"
    ).strip()
    if not header and not entity_counts:
        return RawExtractionResult(
            status=ProcessingJobStatus.UNSUPPORTED,
            parser_name="step_header",
            unsupported_reason="STEP file has no readable ISO-10303 header or entity section.",
        )
    return RawExtractionResult(status=ProcessingJobStatus.COMPLETED, parser_name="step_header", text=text)


def _step_header(content: str) -> str:
    header_start = content.find("HEADER;")
    data_start = content.find("DATA;")
    if header_start == -1:
        return ""
    header_end = data_start if data_start != -1 else min(len(content), header_start + 4000)
    lines = [line.strip() for line in content[header_start:header_end].splitlines()]
    interesting = [
        line
        for line in lines
        if line.startswith(("FILE_DESCRIPTION", "FILE_NAME", "FILE_SCHEMA"))
    ]
    return "\n".join(interesting)


def _step_entity_counts(content: str) -> dict[str, int]:
    counts: dict[str, int] = {}
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("#") or "=" not in line or "(" not in line:
            continue
        entity_name = line.split("=", 1)[1].split("(", 1)[0].strip().upper()
        if not entity_name:
            continue
        counts[entity_name] = counts.get(entity_name, 0) + 1
    return counts


def _create_image_preview(
    document: Document,
    path: Path,
    storage: LocalFileStorage,
) -> list[DocumentArtifactRead]:
    try:
        normalized = _normalize_image(path)
        content = normalized["content"]
        filename = f"{path.stem}.preview.png"
        content_type = "image/png"
    except Exception:
        content = path.read_bytes()
        filename = f"{path.stem}.preview{path.suffix.lower() or '.bin'}"
        content_type = _image_content_type(path)
        normalized = {"width": None, "height": None, "normalized": False}

    storage_path, sha256, size_bytes = storage.save_artifact(document.id, filename, content)
    return [
        DocumentArtifactRead(
            artifact_type="image_preview",
            storage_path=storage_path,
            content_type=content_type,
            width=normalized.get("width"),
            height=normalized.get("height"),
            metadata={
                "source": "image_normalization",
                "normalized": bool(normalized.get("normalized")),
                "sha256": sha256,
                "size_bytes": size_bytes,
            },
        )
    ]


def _create_pdf_page_previews(
    document: Document,
    path: Path,
    storage: LocalFileStorage,
) -> list[DocumentArtifactRead]:
    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError:
        return []

    artifacts: list[DocumentArtifactRead] = []
    try:
        with fitz.open(path) as pdf:
            for page_index in range(min(3, len(pdf))):
                page = pdf[page_index]
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                content = pixmap.tobytes("png")
                filename = f"{path.stem}.page-{page_index + 1}.png"
                storage_path, sha256, size_bytes = storage.save_artifact(
                    document.id,
                    filename,
                    content,
                )
                artifacts.append(
                    DocumentArtifactRead(
                        artifact_type="pdf_page_preview",
                        storage_path=storage_path,
                        content_type="image/png",
                        page_number=page_index + 1,
                        width=pixmap.width,
                        height=pixmap.height,
                        metadata={
                            "source": "pymupdf_render",
                            "sha256": sha256,
                            "size_bytes": size_bytes,
                        },
                    )
                )
    except Exception:
        return []
    return artifacts


def _normalize_image(path: Path) -> dict:
    from PIL import Image  # type: ignore[import-not-found]

    with Image.open(path) as image:
        image = image.convert("RGB")
        image.thumbnail((1600, 1600))
        buffer = BytesIO()
        image.save(buffer, format="PNG", optimize=True)
        return {
            "content": buffer.getvalue(),
            "width": image.width,
            "height": image.height,
            "normalized": True,
        }


def _artifact_as_data_uri(artifact: DocumentArtifactRead) -> str:
    content = Path(artifact.storage_path).read_bytes()
    encoded = base64.b64encode(content).decode("ascii")
    content_type = artifact.content_type or "application/octet-stream"
    return f"data:{content_type};base64,{encoded}"


def _image_content_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".png":
        return "image/png"
    if suffix in {".tif", ".tiff"}:
        return "image/tiff"
    if suffix == ".webp":
        return "image/webp"
    if suffix == ".bmp":
        return "image/bmp"
    return "application/octet-stream"
