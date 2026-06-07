"""Parser registry — map file type to a text extractor.

Every parser takes raw ``bytes`` and returns a :class:`ParsedDocument`. Parsers
never raise: on failure they return empty text with a descriptive
``parser_name`` so the caller can decide whether to fall back to OCR.

Confidentiality note: all parsing is local (no network, no cloud). OCR is left
to the caller; here we only flag ``needs_ocr`` for image / empty-text-layer PDF.
"""

from __future__ import annotations

import io
import json
from dataclasses import dataclass, field
from pathlib import Path

import structlog

logger = structlog.get_logger()


# Extension groups (lowercase, with leading dot).
TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".log"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx", ".xlsm"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif"}
EMAIL_EXTENSIONS = {".eml"}
STEP_EXTENSIONS = {".step", ".stp"}
PDF_EXTENSIONS = {".pdf"}
DXF_EXTENSIONS = {".dxf"}
DWG_EXTENSIONS = {".dwg"}


@dataclass
class ParsedDocument:
    """Result of parsing a document's bytes into text.

    Attributes:
        text: Extracted plain text (empty if none / needs OCR).
        parser_name: Identifier of the parser that ran (for telemetry/UI).
        needs_ocr: True when the input is an image or a PDF with no text layer,
            i.e. the caller should run the VLM OCR fallback.
        page_count: Number of pages for paginated formats (PDF), else None.
        meta: Optional structured metadata (e.g. parsed email headers).
    """

    text: str = ""
    parser_name: str = "unknown"
    needs_ocr: bool = False
    page_count: int | None = None
    meta: dict = field(default_factory=dict)


def _ext_of(filename: str | None) -> str:
    return Path(filename or "").suffix.lower()


def parse_document(
    content: bytes,
    filename: str | None = None,
    mime_type: str | None = None,
) -> ParsedDocument:
    """Extract text from raw document bytes, dispatching by extension/MIME.

    Args:
        content: Raw file bytes.
        filename: Original filename (used for extension dispatch).
        mime_type: Optional MIME type (used as a secondary hint).

    Returns:
        ParsedDocument. Never raises — parser errors are swallowed and surfaced
        via empty text + a descriptive ``parser_name``.
    """
    if not content:
        return ParsedDocument(parser_name="empty_input")

    ext = _ext_of(filename)
    mime = (mime_type or "").lower()

    # PDF — prefer the text layer; flag for OCR when empty (scanned).
    if ext in PDF_EXTENSIONS or mime == "application/pdf":
        return _parse_pdf(content)

    # Office documents.
    if ext in DOCX_EXTENSIONS or "wordprocessingml" in mime:
        return _parse_docx(content)
    if ext in XLSX_EXTENSIONS or "spreadsheetml" in mime:
        return _parse_xlsx(content)

    # Email.
    if ext in EMAIL_EXTENSIONS or mime == "message/rfc822":
        return _parse_eml(content)

    # CAD.
    if ext in DXF_EXTENSIONS:
        return _parse_dxf(content)
    if ext in DWG_EXTENSIONS:
        return _parse_dwg(content)
    if ext in STEP_EXTENSIONS or "step" in mime:
        return _parse_step(content)

    # Images — defer to OCR.
    if ext in IMAGE_EXTENSIONS or mime.startswith("image/"):
        return ParsedDocument(parser_name="image", needs_ocr=True)

    # Plain text (and friends).
    if ext in TEXT_EXTENSIONS or mime.startswith("text/"):
        return _parse_plain_text(content, ext)

    # Unknown extension — try utf-8, but only accept the result if it actually
    # looks like text. Binary formats we cannot parse (legacy .doc, .odt zips,
    # etc.) would otherwise decode into replacement-char garbage; return empty
    # so the caller skips extraction cleanly instead of storing mush.
    decoded = _safe_decode(content)
    if _looks_textual(decoded):
        return ParsedDocument(text=decoded, parser_name=f"fallback_text{ext or ''}")
    return ParsedDocument(parser_name=f"unsupported{ext or ''}", needs_ocr=False)


# --------------------------------------------------------------------------- #
# Individual parsers (lazy imports, never raise).
# --------------------------------------------------------------------------- #


def _safe_decode(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


def _looks_textual(text: str) -> bool:
    """Heuristic: is a decoded string real text rather than binary garbage?

    Rejects strings dominated by Unicode replacement chars (from a failed utf-8
    decode of binary) or NUL/control bytes.
    """
    if not text.strip():
        return False
    sample = text[:4000]
    bad = sum(1 for ch in sample if ch == "�" or (ord(ch) < 9 and ch not in "\t\n\r"))
    return bad / max(len(sample), 1) < 0.05


def _parse_plain_text(content: bytes, ext: str) -> ParsedDocument:
    return ParsedDocument(text=_safe_decode(content), parser_name=f"text{ext or ''}")


def _parse_pdf(content: bytes) -> ParsedDocument:
    """Extract the PDF text layer via PyMuPDF; flag for OCR when empty."""
    try:
        from app.ai.pdf_processor import extract_pdf

        data = extract_pdf(content, render_pages=False)
        text = data.full_text or ""
        if text.strip():
            return ParsedDocument(
                text=text, parser_name="pdf_text_layer", page_count=data.page_count
            )
        # Text layer empty → scanned PDF → OCR fallback.
        return ParsedDocument(
            parser_name="pdf_scanned", needs_ocr=True, page_count=data.page_count
        )
    except Exception as exc:
        logger.warning("pdf_parse_failed", error=str(exc))
        # Cannot read text layer — still let the caller try OCR.
        return ParsedDocument(parser_name="pdf_error", needs_ocr=True)


def _parse_docx(content: bytes) -> ParsedDocument:
    try:
        import docx  # python-docx
    except ImportError:
        logger.warning("docx_parser_unavailable")
        return ParsedDocument(parser_name="docx_unavailable")
    try:
        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in document.paragraphs]
        table_rows: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
        text = "\n".join(part for part in paragraphs + table_rows if part)
        return ParsedDocument(text=text, parser_name="docx")
    except Exception as exc:
        logger.warning("docx_parse_failed", error=str(exc))
        return ParsedDocument(parser_name="docx_error")


def _parse_xlsx(content: bytes) -> ParsedDocument:
    try:
        import openpyxl
    except ImportError:
        logger.warning("xlsx_parser_unavailable")
        return ParsedDocument(parser_name="xlsx_unavailable")
    try:
        workbook = openpyxl.load_workbook(
            io.BytesIO(content), read_only=True, data_only=True
        )
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    lines.append(" | ".join(values))
        workbook.close()
        return ParsedDocument(text="\n".join(lines), parser_name="xlsx")
    except Exception as exc:
        logger.warning("xlsx_parse_failed", error=str(exc))
        return ParsedDocument(parser_name="xlsx_error")


def _parse_eml(content: bytes) -> ParsedDocument:
    """Parse a raw RFC822 email into a readable text block + header metadata.

    Reuses :func:`app.tasks.imap_client.parse_email_message` so direct ``.eml``
    upload and IMAP polling share one MIME parser. Attachment bytes are exposed
    via ``meta['attachments']`` for the ingest layer to store as linked docs.
    """
    try:
        from app.tasks.imap_client import parse_email_message

        parsed = parse_email_message(content)
        header_lines = [
            f"От: {parsed.from_address}",
            f"Кому: {', '.join(parsed.to_addresses)}",
            f"Тема: {parsed.subject}",
        ]
        if parsed.sent_at:
            header_lines.append(f"Дата: {parsed.sent_at.isoformat()}")
        body = parsed.body_text or _strip_html(parsed.body_html)
        text = "\n".join(header_lines) + "\n\n" + body
        return ParsedDocument(
            text=text.strip(),
            parser_name="eml",
            meta={
                "subject": parsed.subject,
                "from": parsed.from_address,
                "message_id": parsed.message_id,
                "attachments": [
                    {
                        "filename": att.filename,
                        "content": att.content,
                        "content_type": att.content_type,
                        "sha256": att.sha256,
                    }
                    for att in parsed.attachments
                ],
            },
        )
    except Exception as exc:
        logger.warning("eml_parse_failed", error=str(exc))
        return ParsedDocument(parser_name="eml_error")


def _strip_html(html: str) -> str:
    if not html:
        return ""
    import re

    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", html, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _parse_dxf(content: bytes) -> ParsedDocument:
    """Extract text annotations (TEXT/MTEXT/DIMENSION/etc.) from a DXF."""
    text = _dxf_text_entities(content)
    return ParsedDocument(text=text, parser_name="dxf")


def _parse_dwg(content: bytes) -> ParsedDocument:
    """Convert DWG → DXF via libredwg, then extract text annotations."""
    import os
    import shutil
    import subprocess
    import tempfile

    with tempfile.TemporaryDirectory(prefix="dwg_text_") as tmpdir:
        dwg_path = os.path.join(tmpdir, "input.dwg")
        dxf_path = os.path.join(tmpdir, "input.dxf")
        with open(dwg_path, "wb") as f:
            f.write(content)

        dwg2dxf_bin = shutil.which("dwg2dxf")
        if dwg2dxf_bin:
            try:
                result = subprocess.run(
                    [dwg2dxf_bin, "--as", "R2018", "-o", dxf_path, dwg_path],
                    timeout=60,
                    capture_output=True,
                )
                if result.returncode == 0 and os.path.exists(dxf_path):
                    with open(dxf_path, "rb") as f:
                        text = _dxf_text_entities(f.read())
                    if text.strip():
                        return ParsedDocument(
                            text=f"[DWG конвертирован в DXF]\n{text}",
                            parser_name="dwg_via_dxf",
                        )
            except Exception as exc:
                logger.warning("dwg_convert_failed", error=str(exc))

    return ParsedDocument(
        text=f"Технический чертёж DWG. Размер файла: {len(content)} байт. Требуется конвертация.",
        parser_name="dwg_stub",
        meta={
            "warning": (
                "Утилита dwg2dxf (libredwg) не найдена — автоматическое извлечение текста "
                "из DWG невозможно. Пожалуйста, конвертируйте файл в DXF или PDF вручную "
                "перед загрузкой, либо установите libredwg: https://savannah.gnu.org/projects/libredwg/"
            )
        },
    )


def _dxf_text_entities(content: bytes) -> str:
    """Extract all text annotations from DXF bytes using ezdxf."""
    try:
        import ezdxf
        import ezdxf.recover as recover
    except ImportError:
        return _safe_decode(content)[:2000]

    doc = None
    try:
        doc = ezdxf.read(io.StringIO(content.decode("utf-8", errors="replace")))
    except Exception:
        pass

    if doc is None:
        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".dxf", delete=False) as tf:
            tf.write(content)
            tmp_path = tf.name
        try:
            try:
                doc = ezdxf.readfile(tmp_path)
            except Exception:
                doc, _ = recover.readfile(tmp_path)
        except Exception as exc:
            logger.warning("dxf_text_extraction_failed", error=str(exc))
            return _safe_decode(content)[:2000]
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    if doc is None:
        return _safe_decode(content)[:2000]

    texts: list[str] = []
    msp = doc.modelspace()
    for entity in msp:
        etype = entity.dxftype()
        try:
            if etype in ("TEXT", "ATTRIB", "ATTDEF"):
                t = str(entity.dxf.text or "")
                if t:
                    texts.append(t)
            elif etype == "MTEXT":
                t = entity.plain_mtext()
                if t:
                    texts.append(t)
            elif etype == "DIMENSION":
                try:
                    m = entity.dxf.actual_measurement
                    if m:
                        texts.append(str(m))
                except Exception:
                    pass
                try:
                    ov = entity.dxf.text
                    if ov:
                        texts.append(ov)
                except Exception:
                    pass
            elif etype == "TOLERANCE":
                try:
                    t = entity.dxf.string or ""
                    if t:
                        texts.append(t)
                except Exception:
                    pass
        except Exception:
            continue

    return "\n".join(texts)


def _parse_step(content: bytes) -> ParsedDocument:
    """Read the ISO-10303 header + entity counts (no geometry backend)."""
    text_content = _safe_decode(content)
    header = _step_header(text_content)
    entity_counts = _step_entity_counts(text_content)
    if not header and not entity_counts:
        return ParsedDocument(parser_name="step_empty")
    text = (
        "STEP header preview. Full geometry analysis requires FreeCAD/pythonOCC backend.\n"
        f"{header}\n"
        f"Entity counts: {json.dumps(entity_counts, ensure_ascii=False, sort_keys=True)}"
    ).strip()
    return ParsedDocument(text=text, parser_name="step_header")


def _step_header(content: str) -> str:
    header_start = content.find("HEADER;")
    data_start = content.find("DATA;")
    if header_start == -1:
        return ""
    header_end = data_start if data_start != -1 else min(len(content), header_start + 4000)
    lines = [line.strip() for line in content[header_start:header_end].splitlines()]
    interesting = [
        line
        for line in lines
        if line.startswith(("FILE_DESCRIPTION", "FILE_NAME", "FILE_SCHEMA"))
    ]
    return "\n".join(interesting)


def _step_entity_counts(content: str) -> dict[str, int]:
    counts: dict[str, int] = {}
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("#") or "=" not in line or "(" not in line:
            continue
        entity_name = line.split("=", 1)[1].split("(", 1)[0].strip().upper()
        if not entity_name:
            continue
        counts[entity_name] = counts.get(entity_name, 0) + 1
    return counts
